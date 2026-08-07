"""Chunks the technical design doc (docs/Banner_AR_POC_Technical_Design.docx) by
its own "Heading 1" sections (1. Executive Summary ... 12. Screenshot of pages),
supplementing the curated poc_help_metadata rows in help_kb.py with the full
document. Parsed directly via python-docx -- no PDF conversion step, so this
has no dependency on MS Word/LibreOffice being installed.

Images are attached as reference metadata only (a path the frontend can show
alongside the answer) -- they are never sent to the LLM as model input.
"""
import os

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

_DOCX_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "Banner_AR_POC_Technical_Design.docx")

# The doc's own section 12 embeds these 3 screenshots in this order (TTVDCAT,
# TSADETC, TSADETL) -- extracted once from the docx's media and committed
# under docs/screenshots/ rather than re-parsed from the docx at request time.
_SCREENSHOTS = ["screenshots/ttvdcat.png", "screenshots/tsadetc.png", "screenshots/tsadetl.png"]


def _iter_block_items(document):
    """Yield paragraphs and tables in document order -- python-docx exposes
    document.paragraphs and document.tables as two separate flat lists, which
    loses the interleaving (a table's surrounding context matters here, e.g.
    section 3's field-mapping tables), so walk the body XML directly instead."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_to_text(table) -> str:
    lines = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def load_design_doc_chunks() -> list:
    """Returns [{heading, text, source, images}] -- one chunk per Heading-1
    section. Returns [] if the doc isn't present (this source is optional;
    the curated help_metadata rows work fine on their own)."""
    if not os.path.isfile(_DOCX_PATH):
        return []

    document = Document(_DOCX_PATH)
    chunks = []
    current_heading = None
    current_parts = []

    def flush():
        if current_heading and current_parts:
            text = "\n".join(p for p in current_parts if p.strip())
            images = _SCREENSHOTS if current_heading.startswith("12.") else []
            chunks.append({"heading": current_heading, "text": text, "source": "design_doc", "images": images})

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            if block.style.name == "Heading 1":
                flush()
                current_heading = block.text.strip()
                current_parts = []
                continue
            if current_heading is None:
                continue  # title/purpose lines before section 1 -- not part of any chunk
            text = block.text.strip()
            if text:
                current_parts.append(text)
        else:  # Table
            if current_heading is None:
                continue
            table_text = _table_to_text(block)
            if table_text:
                current_parts.append(table_text)

    flush()
    return chunks
